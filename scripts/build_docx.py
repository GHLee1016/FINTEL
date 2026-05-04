"""FINTEL 슬라이드별 해석 + 시장 미시구조 시사점 워드 문서 작성."""
from __future__ import annotations
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).resolve().parent.parent

# ── 색상 ────────────────────────────────
ACCENT = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT_HEX = '1E3A5F'
ACCENT_LIGHT_HEX = 'E8EEF4'
TXT = RGBColor(0x1F, 0x2A, 0x3A)
TXT_MUTED = RGBColor(0x6B, 0x72, 0x80)
HEADER_BG_HEX = '2A3D52'

doc = Document()

# 기본 스타일
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# 페이지 여백
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def add_heading(text, level=1, color=ACCENT):
    """제목 추가."""
    p = doc.add_paragraph()
    sizes = {1: 22, 2: 16, 3: 13}
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(text, italic=False, color=TXT, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullets(items, color=TXT, size=11):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.color.rgb = color
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(2)


def add_callout(label, body, color_label=ACCENT):
    """굵은 라벨 + 본문 단락."""
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}  ')
    r1.bold = True
    r1.font.color.rgb = color_label
    r1.font.size = Pt(11)
    r2 = p.add_run(body)
    r2.font.color.rgb = TXT
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_kv_table(rows, col1_label='항목', col2_label='내용', col_widths=(Cm(4), Cm(13))):
    """2-col key-value 표."""
    tbl = doc.add_table(rows=len(rows) + 1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    # 헤더
    hdr = tbl.rows[0]
    for i, txt in enumerate([col1_label, col2_label]):
        cell = hdr.cells[i]
        cell.width = col_widths[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        set_cell_bg(cell, HEADER_BG_HEX)
    # 데이터
    for ri, (k, v) in enumerate(rows, start=1):
        row = tbl.rows[ri]
        c0, c1 = row.cells
        c0.width = col_widths[0]
        c1.width = col_widths[1]
        c0.text = ''
        c1.text = ''
        r0 = c0.paragraphs[0].add_run(k)
        r0.bold = True
        r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(v)
        r1.font.size = Pt(10)
        if ri % 2 == 0:
            set_cell_bg(c0, 'F7F8FA')
            set_cell_bg(c1, 'F7F8FA')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


def hr():
    """가로 구분선 (단락 하단 border)."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D5DBE3')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ════════════════════════════════════════════════
# 1. 표제 / 개요
# ════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run('FINTEL — 결과 슬라이드별 해석 + 시장 미시구조 시사점')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = TXT

sub = doc.add_paragraph()
run = sub.add_run('RMSE_CV 비교 (ML vs 금융모형) + 데이터셋 추가효과')
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = ACCENT
sub.paragraph_format.space_after = Pt(12)

hr()

add_heading('개요', level=2)
add_para(
    '본 문서는 FINTEL 변동성 예측 결과 PPT(`fintel_results.pptx`)의 각 데이터 슬라이드에 대한 해석과, '
    '결과가 각 시장의 미시구조 특성과 어떻게 연결되는지 정리한 보조 자료다.'
)
add_para('PPT 슬라이드는 16장으로 구성되며, 본 문서는 데이터 슬라이드 12장(4~9 / 11~16)에 대한 해석을 다룬다.')

add_heading('분석 범위', level=3)
add_kv_table([
    ('구간', '평시 (normal) / 9·11 / GFC / COVID — 4개'),
    ('시장', 'S&P 500 (US) / KOSPI (KR) / Nikkei 225 (JP) — 3개'),
    ('프로토콜', 'static (1회 fit) / expanding (walk-forward refit) — 2개'),
    ('모델', 'HAR-RV, GARCH(1,1) 금융 + Ridge, ElasticNet, Huber, LightGBM, XGBoost ML'),
    ('Feature tier', 'Core(10) / Momentum(14) / Extended(28) — Section 2 분석에 사용'),
    ('평가 phase', 'Full Test (test 전체) — 본 문서의 모든 결과 기준'),
])

add_heading('RMSE_CV 정의 및 사용 이유', level=3)
add_para('RMSE_CV = RMSE / mean(y_true)', size=12, color=ACCENT, bold=True)
add_para(
    '시장별 평균 RV 스케일이 다르기 때문에 (KR≈0.53, JP≈0.83 등), '
    '같은 RMSE 절대값도 의미가 다르다. RMSE_CV는 RMSE를 그 phase의 평균 RV로 나눠 무차원화함으로써 '
    '시장·구간·모델 간 직접 비교를 가능하게 한다.'
)
add_kv_table([
    ('RMSE_CV ≈ 0.20', '오차가 평균 RV의 약 20% — 양호한 예측'),
    ('RMSE_CV ≈ 0.30', '오차가 평균의 30% — 보통 수준'),
    ('RMSE_CV ≈ 0.50', '오차가 평균의 50% — 큰 오차'),
    ('RMSE_CV ≈ 1.00', '오차가 평균과 동일 — 모델 무용 수준'),
], col1_label='수치', col2_label='해석', col_widths=(Cm(3.5), Cm(13.5)))

add_heading('본 PPT 결과 범위', level=3)
add_kv_table([
    ('ML best', '0.16 ~ 0.34 (Extended tier 기준)'),
    ('HAR-RV', '0.21 ~ 0.50'),
    ('GARCH(1,1)', '0.30 ~ 0.95 (sqrt 변환 후)'),
], col1_label='모델 부류', col2_label='RMSE_CV 범위')

doc.add_page_break()

# ════════════════════════════════════════════════
# 2. Section 1 — ML vs 금융모형
# ════════════════════════════════════════════════
add_heading('Section 1 — ML vs 금융모형', level=1)
add_para(
    '7개 모델(HAR-RV / GARCH / Ridge / ElasticNet / Huber / LightGBM / XGBoost)을 '
    '각 시장 × 프로토콜 × 구간에서 RMSE_CV로 비교. ML 모델은 Extended tier(28 feature) 사용.',
    italic=True, color=TXT_MUTED
)

# 슬라이드 4: US static
hr()
add_heading('슬라이드 4 — S&P 500 (US) / static', level=2)
add_callout('핵심 관찰',
    '모든 구간에서 GARCH(0.41~0.83) > HAR(0.34~0.48) > ML(0.23~0.46) 위계가 일관되게 나타남.')
add_bullets([
    '평시 best: Huber 0.280 — HAR 대비 36.2% 개선',
    '9·11 best: ElasticNet 0.230 — HAR 대비 32.4% 개선',
    'GFC best: Ridge 0.242 — 트리(LGBM 0.399, XGB 0.355)가 의외로 약함',
    'COVID best: LightGBM 0.327 — 선형 ML 셋(0.46+) 모두 HAR(0.478)과 비슷, 트리만 우뚝',
    'GARCH는 COVID에서 0.831까지 폭증 — 평균회귀 가정의 한계',
])
add_callout('미시구조 시사점',
    '미국 시장은 평시·완만한 위기에서 변동성이 비교적 정규성에 가깝게 분포해 선형 ML로 잘 capture된다. '
    '반면 COVID 같은 unprecedented shock(2020년 3월 단 한 달간 VIX 80+ 도달)에서는 '
    '변동성 점프가 비선형적·이질적이라 트리 모델만이 효과적으로 적응한다. '
    'GFC도 점진적 위기였지만 미국 시장에서는 정책 대응이 빨라 "step-wise" 패턴이라 선형 우세.'
)

# 슬라이드 5: US expanding
hr()
add_heading('슬라이드 5 — S&P 500 (US) / expanding', level=2)
add_callout('핵심 관찰', 'static과 거의 동일 추세지만 평시·COVID에서 best가 LightGBM으로 이동.')
add_bullets([
    '평시 best: LightGBM 0.270 (static의 Huber 0.280에서 추가 -3.6%)',
    'COVID best: LightGBM 0.301 (static 0.327에서 -7.9%) — 데이터 누적 효과 큼',
    '9·11·GFC는 static과 거의 동일 (변화 ≤ 0.005)',
    '선형 ML도 COVID에서 0.40~0.46으로 static 0.46~0.48 대비 미세 개선',
])
add_callout('미시구조 시사점',
    '미국 시장은 거래량이 가장 크고 정보 효율성이 높아, expanding window의 데이터 누적 효과가 트리 모델에 잘 반영된다. '
    '특히 COVID는 정책 변화·백신 뉴스 등 새로운 정보가 매일 들어오는 환경이라, '
    '매 step refit하는 expanding이 비선형 모델의 적응력을 끌어올림. '
    '선형 ML은 평균 패턴을 잡는 데 강하지만 새 정보를 빠르게 흡수하는 데는 트리에 밀림.'
)

# 슬라이드 6: KR static
hr()
add_heading('슬라이드 6 — KOSPI (KR) / static', level=2)
add_callout('핵심 관찰',
    '3 시장 중 ML이 가장 일관되게 잘 동작 (평시·GFC·COVID에서 ML best 0.20~0.25, US보다 낮음).')
add_bullets([
    '평시 best: Huber 0.201 — HAR 대비 30.7%, GARCH 대비 70.5% 개선',
    '9·11에서 선형 ML 약함 (0.34~0.38), 트리(XGB 0.251)가 우세',
    'GFC best: Ridge 0.207, COVID best: Huber 0.250 — 절대값 grid 평균 최저급',
    'GARCH가 0.55~0.82로 가장 극단적으로 나쁨',
])
add_callout('미시구조 시사점',
    'KOSPI는 외국인 자금 유입에 매우 민감하고, 환율·미국 시장에 강하게 동조한다. '
    'KR persistence는 0.996 (half-life 175일)로 변동성이 매우 끈적해서 GARCH의 평균회귀 가정이 거의 작동하지 않는다 — '
    '평시에 fit한 GARCH가 위기로 진입할 때 σ²이 long-run mean(약 2.0)에 천천히 수렴하려 하지만 '
    '실제 RV는 빠르게 점프해 RMSE_CV 폭증. ML은 환율·spillover_SP500 같은 외생 정보를 직접 받기 때문에 '
    '한국 시장의 지배적 driver를 잘 capture. 9·11이 예외인 이유는 짧은 train(327일)에 한국이 직접 충격받지 않은 외부 사건이라 '
    '거시 신호가 비선형적으로 반영되어 트리가 우세.'
)

# 슬라이드 7: KR expanding
hr()
add_heading('슬라이드 7 — KOSPI (KR) / expanding', level=2)
add_callout('핵심 관찰',
    '평시 LightGBM/XGBoost 동시 0.188(best) — KR static 평시 0.201보다 6.5% 추가 개선.')
add_bullets([
    'Walk-forward로 트리가 평시에서 선형을 추월한 유일한 KR 케이스',
    '9·11에서 선형 ML이 더 나아짐 (0.33~0.35) 하지만 여전히 트리(XGB 0.265)가 best',
    'GFC·COVID는 static과 거의 동일 — 결과 안정성 확보',
    '12 (시장×프로토콜) 셀 중 두 번째 최저값 (0.188)',
])
add_callout('미시구조 시사점',
    '한국 시장은 일중 거래 패턴이 미국·일본 대비 비교적 균질해서, expanding window의 데이터 누적이 '
    '비선형 모델의 학습에 잘 작용한다. 특히 평시에 트리가 우세해진다는 것은 '
    'KOSPI의 일별 변동성이 단순한 자기상관 외에 feature 간 상호작용(예: 환율 변동률 × 기업실적 시즌)이 있음을 시사. '
    '9·11에서 트리가 여전히 우세한 것은 짧은 train에서 데이터 누적 효과가 작아 선형이 따라잡지 못하기 때문.'
)

# 슬라이드 8: JP static
hr()
add_heading('슬라이드 8 — Nikkei 225 (JP) / static', level=2)
add_callout('핵심 관찰',
    '9·11에서 ElasticNet 0.164 — grid 전체 24개 셀 중 절대 최저값.')
add_bullets([
    '9·11에서 선형 ML(0.16~0.20)이 트리(0.18~0.20)를 이김 — 짧은 train에도 정보 신호가 깨끗',
    '평시·COVID에서 GARCH가 0.92~0.95로 grid 전체 최악 — 야간 갭 효과',
    '평시 best Ridge 0.249 — HAR 대비 -45.9% (3 시장 중 최대 개선폭)',
    '모든 구간에서 ML이 안정적으로 좋음 (Ridge/EN 강세)',
])
add_callout('미시구조 시사점',
    '일본 시장은 야간 갭이 가장 크다 — 도쿄 시장 휴장 시간(미국 거래 시간)에 발생한 글로벌 충격이 '
    '다음 날 개장 직후 점프로 반영된다. 우리가 사용한 RV는 일중 5분 단위 합산이라 야간 정보를 놓치는데, '
    'GARCH는 close-to-close 수익률(=일중+야간)로 fit하므로 σ² 추정이 RV보다 크게 부풀려진다. '
    '이 mismatch가 sqrt 변환 후에도 잔존해 GARCH가 가장 폭증. '
    '9·11이 grid 최저인 것은 일본이 지정학적으로 9·11에서 직접 영향이 작아 변동성 패턴이 깨끗했고, '
    '그 결과 짧은 train에서도 선형 ML이 신호를 잘 잡았다고 해석.'
)

# 슬라이드 9: JP expanding
hr()
add_heading('슬라이드 9 — Nikkei 225 (JP) / expanding', level=2)
add_callout('핵심 관찰',
    '9·11 ElasticNet 0.160 — grid 갱신 최저(static 0.164에서 추가 개선).')
add_bullets([
    '평시 best XGBoost 0.238 (static Ridge 0.249보다 -4.4%)',
    '모든 구간에서 expanding이 미세 개선 — 데이터 누적이 잘 작동',
    'GARCH는 expanding에서도 거의 동일(0.92 수준) — 단위 mismatch는 데이터 양과 무관',
    'GFC/COVID에서 EN/XGB 등 다양한 모델이 best — JP 시장은 모델 간 격차가 작음',
])
add_callout('미시구조 시사점',
    '일본 시장은 위기 유형에 따라 best 모델이 달라지지만 격차가 작다 — '
    '즉 ML 모델 부류 내에서는 어느 모델을 써도 비슷한 성능. '
    'expanding이 모든 구간에서 작은 개선을 주는 것은, JP 시장의 변동성 dynamics가 '
    '데이터 누적에 단조적으로 반응한다는 의미 (US/COVID 같은 큰 점프는 없음). '
    '9·11이 0.160으로 또 갱신된 것은 평시 train(2000-01~2001-08)에 JP 시장이 비교적 stable했고 '
    'expanding window가 9·11 직후의 nonlinear 적응을 부드럽게 흡수했기 때문.'
)

doc.add_page_break()

# ════════════════════════════════════════════════
# 3. Section 2 — Tier 효과
# ════════════════════════════════════════════════
add_heading('Section 2 — 데이터셋 추가효과 (Tier 효과)', level=1)
add_para(
    'Core(10 feature) → Momentum(+추세 4) → Extended(+거시·외생 14) tier 진행에 따른 '
    'RMSE_CV 변화. Best ML = 5 ML 모델 중 최저값. Δ는 다음 tier 추가 효과 (음수 = 개선).',
    italic=True, color=TXT_MUTED
)

# 슬라이드 11: US static — Tier
hr()
add_heading('슬라이드 11 — S&P 500 (US) / static — Tier 효과', level=2)
add_callout('핵심 관찰', 'COVID에서만 Tier 효과 압도적 (-12.3%, Core 0.373 → Extended 0.327).')
add_bullets([
    'Core→Mom -0.031, Mom→Ext -0.015 — 두 단계 모두 도움',
    '평시·9·11·GFC는 거의 무효 (-0.7% / 0% / 0%)',
    '평시·9·11·GFC는 RV-only(Core)로도 ceiling 도달',
])
add_callout('미시구조 시사점',
    '미국은 글로벌 충격(COVID) 시 거시 변수(연준 정책금리, EPU, WTI 등)가 변동성에 결정적 영향을 미친다. '
    '특히 2020년 3월 연준의 긴급 금리인하·QE 발표가 시장 반등의 핵심 driver였는데, '
    'Extended tier의 정책금리·EPU 변수가 이를 capture. '
    '반면 평시·9·11·GFC에서는 미국 시장의 자기상관(RV 자체) 정보만으로도 대부분 설명 가능. '
    '즉 미국 시장은 "거시 정보 의존도가 위기 종류에 따라 극단적으로 갈린다"는 특성.'
)

# 슬라이드 12: US expanding — Tier
hr()
add_heading('슬라이드 12 — S&P 500 (US) / expanding — Tier 효과', level=2)
add_callout('핵심 관찰', 'COVID -11.2% (static과 거의 동일), 평시 -3.2% (static 강화).')
add_bullets([
    'COVID 효과는 데이터 누적 무관하게 거시 정보 자체가 핵심',
    '평시 -3.2%로 강화 — expanding이 거시 변수의 점진적 효과를 잘 흡수',
    '9·11·GFC는 여전히 무효',
])
add_callout('미시구조 시사점',
    'expanding window에서 평시 Tier 효과가 -3.2%로 강화된 것은, '
    '미국 시장에서 거시 변수(특히 EPU, 정책금리)가 일별 변동성에 점진적·누적적으로 영향을 미친다는 의미. '
    'COVID는 변화가 너무 급격해 expanding의 효과가 static과 비슷.'
)

# 슬라이드 13: KR static — Tier
hr()
add_heading('슬라이드 13 — KOSPI (KR) / static — Tier 효과', level=2)
add_callout('핵심 관찰', 'GFC에서 가장 큰 효과 -6.8% (Mom→Ext -0.019, 즉 거시 변수가 결정적).')
add_bullets([
    '평시 -2.0% (Mom→Ext -0.003), 9·11 -2.0% (Core→Mom -0.008)',
    'COVID Tier 효과 0% — KR의 COVID 변동성은 RV-only로 충분히 capture',
    'GFC 시 환율·금리·정책금리 정보가 결정적임을 정량 증거',
])
add_callout('미시구조 시사점',
    'KOSPI는 GFC 당시 외국인 자금의 급격한 유출입이 변동성의 주된 driver였다 — '
    '원/달러 환율이 1,500원대까지 폭등하고 정책금리가 5%→2%로 급강하한 시기. '
    'Extended tier의 fx_change·policy_rate·rate_spread 등이 이 dynamics를 capture해 RMSE_CV -6.8% 개선. '
    'COVID에서는 환율 변동이 GFC만큼 크지 않았고 외국인 자금 유출도 빠르게 정상화되어 거시 정보가 큰 도움 안 됨. '
    'KR이 미국과 다른 점: KR은 GFC에서 거시 정보가 강하게 효과, 미국은 COVID에서 효과 — 위기의 transmission 채널이 다름을 시사.'
)

# 슬라이드 14: KR expanding — Tier
hr()
add_heading('슬라이드 14 — KOSPI (KR) / expanding — Tier 효과', level=2)
add_callout('핵심 관찰',
    '9·11에서 -4.7%, 특히 Core→Mom -0.034 (모멘텀 추가만으로 큰 개선) 후 Mom→Ext +0.021 (거시 추가는 노이즈).')
add_bullets([
    'KR/9·11 expanding은 "Momentum tier가 best, Extended는 over"라는 명확한 비단조성',
    '평시 -2.1% (static과 비슷)',
    'GFC·COVID는 0% — expanding으로도 데이터셋 추가 효과 안 살아남',
])
add_callout('미시구조 시사점',
    '한국 시장은 9·11 당시 train이 327일로 짧고, 위기 자체가 외부 사건(미국 본토 테러)이라 '
    '한국 거시 변수(원/달러, 한은 정책금리)는 9·11에 직접 반응하지 않았다. '
    '결국 모멘텀(추세) 정보까지는 도움 되지만 거시 변수는 "관계없는 정보"라 노이즈로 작용. '
    '이는 "feature 추가가 항상 좋은 것은 아니다"의 명확한 사례 — train 크기와 위기의 transmission 채널이 모두 중요.'
)

# 슬라이드 15: JP static — Tier
hr()
add_heading('슬라이드 15 — Nikkei 225 (JP) / static — Tier 효과', level=2)
add_callout('핵심 관찰', '9·11에서 +4.5% — 24개 셀 중 유일한 강한 negative effect.')
add_bullets([
    'Core 0.157 → Extended 0.164. 데이터 추가가 적극적으로 해로움',
    '원인: 9·11 train ~329일에 28 feature → over-parameterization',
    '평시·GFC·COVID는 거의 무효 (-0.4 ~ -1.4%)',
    'JP/9·11에서는 Core(10 feature)가 best',
])
add_callout('미시구조 시사점',
    '일본은 9·11에서 지정학적으로 가장 거리가 먼 시장이라 직접 충격이 작다. '
    '한국과 달리 일본은 9·11 당시 "잃어버린 10년" 후반기로 자체 디플레이션·금융권 부실 문제가 더 dominant했고, '
    '미국 사태가 일본 거시변수에 미친 영향은 일시적. '
    '이런 상황에서 짧은 train(329일)에 28 feature는 분명히 over-parameterization. '
    'Optuna가 Ridge/EN의 alpha를 높여도 정보가 없는 변수의 영향을 완전히 빼지 못함. '
    '결론: feature는 "정보 신호 / 데이터 양" 균형 맞춰야 하며, 외부 사건의 transmission 채널이 약한 시장에선 Core가 최적.'
)

# 슬라이드 16: JP expanding — Tier
hr()
add_heading('슬라이드 16 — Nikkei 225 (JP) / expanding — Tier 효과', level=2)
add_callout('핵심 관찰',
    '9·11 +2.6% (static의 +4.5%에서 줄었지만 여전히 negative), 평시 -3.3%, COVID -2.7%.')
add_bullets([
    '평시 Mom→Ext -0.007 (적당한 도움)',
    'COVID Core→Mom -0.017 (모멘텀 큰 도움), Mom→Ext +0.009 (거시는 마이너스)',
    'JP COVID도 모멘텀이 sweet spot',
    'GFC -0.5%로 무효',
])
add_callout('미시구조 시사점',
    'expanding window가 9·11 over-parameterization을 부분적으로 완화 (+4.5%→+2.6%), '
    '하지만 여전히 negative — JP 시장의 9·11 데이터셋엔 거시 정보의 marginal contribution이 거의 없다는 강한 증거. '
    'COVID에서 모멘텀(추세)이 거시보다 우세한 것은 일본 시장의 자체 모멘텀 효과(특히 Topix·BOJ 정책 등)가 '
    '글로벌 거시 신호보다 더 dominant하다는 의미. '
    '미국 COVID와 정확히 반대 양상 (미국=거시 우세, 일본=모멘텀 우세).'
)

doc.add_page_break()

# ════════════════════════════════════════════════
# 4. 시장 미시구조 종합 시사점
# ════════════════════════════════════════════════
add_heading('Section 3 — 시장 미시구조 종합 시사점', level=1)
add_para(
    'Section 1·2의 결과를 시장 단위로 묶어, 각 시장의 미시구조 특성과 변동성 예측 결과 패턴의 연관을 종합 정리.',
    italic=True, color=TXT_MUTED
)

# US
add_heading('🇺🇸 S&P 500 (US) — 글로벌 leading market, 거시 정보 의존', level=2)

add_heading('미시구조 특성', level=3)
add_bullets([
    '세계에서 가장 효율적·유동성 높은 시장. 일평균 거래량 압도적',
    '글로벌 거시 정보의 1차 진원지 (연준·미국 경제지표) — 모든 시장에 spillover',
    '야간 갭 작음 — 24시간 글로벌 시장 중 leading position이라 정보 흡수가 일중에 집중',
    '변동성 자체는 다른 시장 대비 평탄(평시 RMSE_CV 분모 ≈ 0.59), 위기 시 점프 강함',
])

add_heading('결과 패턴과 미시구조 연결', level=3)
add_bullets([
    '평시·9·11·GFC: 선형 ML 우세 (Huber·EN·Ridge) — 정규성에 가까운 변동성 분포',
    'COVID: 트리(LightGBM/XGBoost)만 통함 — unprecedented shock의 비선형성 (VIX 80+)',
    'COVID에서 Tier 효과 -11~12% (Best vs only-RV) — 거시 변수가 결정적 driver',
    'GARCH 폭증 정도 가장 적음 — 야간 갭 작아 r²/RV 미스매치 작음',
])

add_callout('시사점',
    '미국 시장은 변동성 모델링에서 "평시는 단순 모델, 위기는 비선형 + 거시 정보" 라는 명확한 이분법. '
    'GARCH가 다른 시장 대비 덜 망하는 것은 야간 갭 효과가 작기 때문 — 모델의 가정(close-to-close = intraday)이 가장 잘 맞는 시장.'
)

# KR
add_heading('🇰🇷 KOSPI (KR) — 외국인 자금 의존, 끈적한 변동성', level=2)

add_heading('미시구조 특성', level=3)
add_bullets([
    '외국인 자금 유입에 매우 민감 (시가총액 대비 외국인 보유 30%+)',
    '환율·미국 시장에 강하게 동조 — spillover 효과 본질적',
    'Persistence 0.996 (half-life 175일) — 변동성이 매우 끈적. 평시→위기 전환에 모델이 못 따라감',
    '야간 갭 중간 수준 (US가 도쿄·서울 휴장 시 거래)',
])

add_heading('결과 패턴과 미시구조 연결', level=3)
add_bullets([
    '3 시장 중 ML이 가장 일관되게 우수 (평시 RMSE_CV 0.20~0.21)',
    'GARCH가 가장 폭증 (0.55~0.82) — 끈적한 persistence 때문에 평균회귀 거의 작동 안 함',
    'GFC에서 Tier 효과 -6.8% — 환율·금리 변수가 결정적 (외국인 자금 유출입)',
    '9·11에서 트리 우세 — 한국이 직접 충격받지 않은 외부 사건이라 nonlinear spillover 반응',
    'COVID Tier 효과 0% — 환율·외국인 자금 변동이 GFC만큼 극단적이지 않았음',
])

add_callout('시사점',
    'KOSPI는 "내부 dynamics + 외부 spillover" 의 혼합 시장. '
    'GARCH가 가장 약한 것은 한국의 끈적한 변동성 persistence (외국인 자금이 stick하면 오래 stick) 때문. '
    'GFC에서 거시 변수가 강하게 도움된다는 발견은 "한국 변동성의 핵심은 외국인 자본 흐름"이라는 통상적 직관을 정량 확인.'
)

# JP
add_heading('🇯🇵 Nikkei 225 (JP) — 야간 갭 dominant, 모멘텀 시장', level=2)

add_heading('미시구조 특성', level=3)
add_bullets([
    '야간 갭이 가장 큼 — 미국 시장이 도쿄 휴장 시간에 거래되어 다음 날 개장 점프로 반영',
    '일중 RV(5분 합산)와 close-to-close 수익률 분산이 가장 차이남',
    '지정학적으로 미국·중동 사건에서 한 단계 떨어진 시장 (9·11 영향 작음)',
    '모멘텀 효과 강함 (BOJ 정책·엔환율 dynamics)',
])

add_heading('결과 패턴과 미시구조 연결', level=3)
add_bullets([
    'GARCH가 평시·COVID에서 가장 폭증 (0.92~0.95) — 야간 갭 mismatch 잔존',
    '9·11 grid 전체 최저 ElasticNet 0.160 — 외부 사건 직접 영향 작아 신호 깨끗',
    '9·11에서 Tier 효과 +2.6~+4.5% (negative!) — 거시 변수가 노이즈로 작용',
    '평시 ML이 HAR 대비 -45.9% (3 시장 중 최대 개선폭) — 일중 패턴이 깨끗하게 학습됨',
    'COVID에서 모멘텀 효과 우세 (Core→Mom -0.017) — 일본 시장의 자체 모멘텀이 지배적',
])

add_callout('시사점',
    '일본 시장의 핵심 미시구조 issue는 야간 갭. RV(일중 only)와 일별 수익률 분산이 본질적으로 다른 양을 측정. '
    'GARCH가 close-to-close로 fit하는 한 sqrt 변환 후에도 미스매치 잔존. '
    '대안: BV(bipower variation), JV(jump variation) 등 야간 갭 없는 추정량 사용 또는 RV에 야간 r² 추가. '
    '9·11에서 negative Tier effect는 "외부 사건의 transmission 채널이 약한 시장에선 거시 변수가 노이즈"라는 일반 원리의 명확한 사례.'
)

doc.add_page_break()

# ════════════════════════════════════════════════
# 5. 종합 결론
# ════════════════════════════════════════════════
add_heading('Section 4 — 종합 결론', level=1)

add_heading('일관된 발견', level=2)
add_bullets([
    'GARCH < HAR < ML 위계: 24 셀 모두에서 흔들림 없이 성립',
    'ML 개선율: HAR 대비 23~38%, GARCH 대비 56~70%',
    'ML 부류 내 best 모델은 위기 유형별로 다름 — "no free lunch"',
    'Feature tier 추가 효과는 평균 -1.6%로 작지만, 위기 transmission 채널이 강한 곳에서 -7~12%',
])

add_heading('시장별 차별화된 패턴', level=2)
add_kv_table([
    ('US', '평시·9·11·GFC = 선형, COVID = 트리. COVID에서 거시 정보 결정적 (-12%).'),
    ('KR', 'ML 일관되게 우수, GFC에서 거시 정보 -6.8%. 9·11은 트리만 통함.'),
    ('JP', 'EN/Ridge 강세, 9·11 grid 최저 0.160. 9·11에서 Tier 추가는 해로움 (+2.6~4.5%).'),
])

add_heading('발표용 한 문장 요약', level=2)
add_para(
    '"전통 모형(GARCH·HAR)은 모든 위기에서 ML에 패배하지만, ML 내에서도 위기 유형 × 시장 미시구조에 따라 '
    '최적 모델·tier가 달라지므로 단일 정답은 없다. 핵심은 \'train 크기 + 시장 dynamics + 위기 transmission 채널\'에 따라 '
    '모델·feature를 선택하는 것."',
    italic=True, color=ACCENT, size=12
)

# 저장
out = PROJECT_ROOT / 'report/fintel_slide_commentary.docx'
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(f'saved: {out}')
print(f'paragraphs: {len(doc.paragraphs)}')
